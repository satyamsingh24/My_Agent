"""Core CV storage, PDF extraction, JD matching, and ATS PDF generation."""
from __future__ import annotations

import html
import json
import re
import secrets
from datetime import datetime
from functools import lru_cache
from io import BytesIO
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from openpyxl import Workbook
from openpyxl.drawing.image import Image as XLImage
from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
from pypdf import PdfReader
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import Image as RLImage
from reportlab.platypus import (
    HRFlowable,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)
from PIL import Image as PILImage
from PIL import ImageDraw, ImageFont

ROOT = Path(__file__).resolve().parent
DATA_DIR = ROOT / "app_data"
CV_PDF = DATA_DIR / "existing_cv.pdf"
CV_TEXT = DATA_DIR / "existing_cv.txt"
CV_META = DATA_DIR / "existing_cv.json"
OUTPUT_DIR = ROOT / "applications" / "generated"

CV_TEMPLATES = {
    "reference": {
        "label": "Reference Black & White",
        "description": (
            "Exact layout of the supplied reference CV: no colour, ruled "
            "headings, and tight spacing."
        ),
        "primary": "000000",
        "accent": "000000",
        "heading_fill": "FFFFFF",
        "heading_text": "000000",
        "ruled_headings": True,
        "compact": True,
    },
    "modern": {
        "label": "Modern Blue",
        "description": "A fresh blue design with softly highlighted section headings.",
        "primary": "155E75",
        "accent": "0E7490",
        "heading_fill": "DDF3F7",
        "heading_text": "155E75",
        "ruled_headings": False,
        "compact": False,
    },
    "minimal": {
        "label": "Clean Minimal",
        "description": "A restrained charcoal design for a simple corporate CV.",
        "primary": "20242C",
        "accent": "4B5563",
        "heading_fill": "E7E9ED",
        "heading_text": "20242C",
        "ruled_headings": False,
        "compact": False,
    },
    "classic_navy": {
        "label": "Classic Navy",
        "description": "Traditional navy headings for finance, consulting, and corporate roles.",
        "primary": "172554",
        "accent": "1E3A8A",
        "heading_fill": "E8EEF9",
        "heading_text": "172554",
        "ruled_headings": True,
        "compact": False,
    },
    "executive_maroon": {
        "label": "Executive Maroon",
        "description": "A confident maroon design suited to senior and leadership profiles.",
        "primary": "701A2C",
        "accent": "9F1239",
        "heading_fill": "FCE7EC",
        "heading_text": "701A2C",
        "ruled_headings": False,
        "compact": True,
    },
    "forest": {
        "label": "Forest Professional",
        "description": "Calm forest-green accents with clear, compact section bands.",
        "primary": "14532D",
        "accent": "15803D",
        "heading_fill": "E5F5EA",
        "heading_text": "14532D",
        "ruled_headings": False,
        "compact": True,
    },
    "slate": {
        "label": "Slate Technical",
        "description": "Sharp slate styling for engineering, DevOps, and technical roles.",
        "primary": "1E293B",
        "accent": "475569",
        "heading_fill": "E2E8F0",
        "heading_text": "1E293B",
        "ruled_headings": True,
        "compact": True,
    },
    "royal_purple": {
        "label": "Royal Purple",
        "description": "A polished purple layout for product, design, and technology roles.",
        "primary": "581C87",
        "accent": "7E22CE",
        "heading_fill": "F3E8FF",
        "heading_text": "581C87",
        "ruled_headings": False,
        "compact": False,
    },
    "teal_compact": {
        "label": "Teal Compact",
        "description": "Space-efficient teal styling for candidates with detailed experience.",
        "primary": "134E4A",
        "accent": "0F766E",
        "heading_fill": "DDF4F1",
        "heading_text": "134E4A",
        "ruled_headings": True,
        "compact": True,
    },
    "warm_gray": {
        "label": "Warm Gray",
        "description": "A neutral, understated format suitable for every professional domain.",
        "primary": "292524",
        "accent": "57534E",
        "heading_fill": "EDEAE7",
        "heading_text": "292524",
        "ruled_headings": False,
        "compact": False,
    },
}


def cv_template(template: str) -> dict[str, str]:
    """Return a supported CV template, falling back to the reference design."""
    return CV_TEMPLATES.get(template, CV_TEMPLATES["reference"])


# Preview thumbnails try a real font first, then Pillow's built-in font so the
# picker also works in the browser build where system fonts are missing.
PREVIEW_FONTS = {
    False: (
        "C:/Windows/Fonts/arial.ttf",
        "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
        "/Library/Fonts/Arial.ttf",
    ),
    True: (
        "C:/Windows/Fonts/arialbd.ttf",
        "/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf",
        "/Library/Fonts/Arial Bold.ttf",
    ),
}


@lru_cache(maxsize=16)
def _preview_font(size: int, bold: bool = False):
    for path in PREVIEW_FONTS[bold]:
        try:
            return ImageFont.truetype(path, size)
        except OSError:
            continue
    try:
        return ImageFont.load_default(size=size)
    except TypeError:
        return ImageFont.load_default()


@lru_cache(maxsize=16)
def _font_at(path: str, size: int):
    return ImageFont.truetype(path, size)


@lru_cache(maxsize=4096)
def _cached_width(text: str, path: str, size: int) -> float:
    return _font_at(path, size).getlength(text)


@lru_cache(maxsize=16)
def _wide_char_width(path: str, size: int) -> float:
    """Width of the widest common character, used to skip exact measuring."""
    font = _font_at(path, size)
    return max(font.getlength(char) for char in "WM@%_")


def _text_width(text: str, font) -> float:
    """Glyph width of `text`, cached because font measuring is slow."""
    path = getattr(font, "path", None)
    size = getattr(font, "size", None)
    if not isinstance(path, str) or not isinstance(size, int):
        return font.getlength(text)
    return _cached_width(text, path, size)


def _certainly_fits(text: str, font, limit: float) -> bool:
    """True when even all-wide glyphs would fit, so no measuring is needed."""
    path = getattr(font, "path", None)
    size = getattr(font, "size", None)
    if not isinstance(path, str) or not isinstance(size, int):
        return False
    return len(text) * _wide_char_width(path, size) <= limit


def _fit_preview_text(text: str, font, limit: float) -> str:
    """Shorten text to `limit` pixels using few (slow) width measurements."""
    if _certainly_fits(text, font, limit):
        return text
    width = _text_width(text, font)
    if width <= limit:
        return text
    keep = max(0, min(len(text) - 1, int(len(text) * limit / width) - 1))
    while keep and _text_width(text[:keep] + "...", font) > limit:
        keep -= max(1, keep // 8)
    while keep < len(text) - 1 and _text_width(text[:keep + 1] + "...", font) <= limit:
        keep += 1
    return (text[:keep] + "...") if keep else ""


def build_template_preview(text: str, template: str, width: int = 430) -> bytes:
    """Render a small page thumbnail so a template can be judged before use."""
    design = cv_template(template)
    height = int(width * 1.414)
    image = PILImage.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(image)
    draw.rectangle([(0, 0), (width - 1, height - 1)], outline="#D8DCE8")

    primary = "#" + design["primary"]
    accent = "#" + design["accent"]
    heading_fill = "#" + design["heading_fill"]
    heading_text = "#" + design["heading_text"]

    margin = int(width * 0.06)
    inner = width - 2 * margin
    name_font = _preview_font(max(13, int(width * 0.048)), bold=True)
    role_font = _preview_font(max(9, int(width * 0.030)), bold=True)
    heading_font = _preview_font(max(8, int(width * 0.026)), bold=True)
    entry_font = _preview_font(max(7, int(width * 0.024)), bold=True)
    body_font = _preview_font(max(7, int(width * 0.023)))

    y = margin
    section = ""
    nonempty_seen = 0
    for raw in text.splitlines():
        line = raw.strip()
        if y > height - margin:
            break
        if not line:
            y += 4
            continue
        upper = line.upper().rstrip(":")
        if nonempty_seen == 0:
            body = _fit_preview_text(line, name_font, inner)
            draw.text((margin, y), body, font=name_font, fill=primary)
            y += name_font.size + 6
        elif nonempty_seen == 1 and line == upper and len(line.split()) <= 6:
            body = _fit_preview_text(line, role_font, inner)
            draw.text((margin, y), body, font=role_font, fill=accent)
            y += role_font.size + 5
        elif upper in SECTION_HEADINGS:
            section = HEADING_ALIASES.get(upper, upper)
            if design.get("ruled_headings"):
                draw.text(
                    (margin, y),
                    _fit_preview_text(upper, heading_font, inner),
                    font=heading_font,
                    fill=heading_text,
                )
                y += heading_font.size + 3
                draw.line(
                    [(margin, y), (width - margin, y)], fill="#000000", width=1
                )
                y += 4
            else:
                band = heading_font.size + 8
                draw.rectangle(
                    [(margin, y), (width - margin, y + band)], fill=heading_fill
                )
                draw.text(
                    (margin + 5, y + 4),
                    _fit_preview_text(upper, heading_font, inner - 10),
                    font=heading_font,
                    fill=heading_text,
                )
                y += band + 5
        elif line.startswith(("-", "*", "\u2022")) or _is_numbered_item(line):
            body = line
            if not _is_numbered_item(line):
                body = "\u2022 " + line.lstrip("-*\u2022 ").strip()
            draw.text(
                (margin + 6, y),
                _fit_preview_text(body, body_font, inner - 6),
                font=body_font,
                fill="#4A5060",
            )
            y += body_font.size + 4
        elif section and section not in PROSE_SECTIONS:
            draw.text(
                (margin, y),
                _fit_preview_text(line, entry_font, inner),
                font=entry_font,
                fill="#1F2437",
            )
            y += entry_font.size + 4
        else:
            draw.text(
                (margin, y),
                _fit_preview_text(line, body_font, inner),
                font=body_font,
                fill="#4A5060",
            )
            y += body_font.size + 4
        nonempty_seen += 1

    output = BytesIO()
    image.save(output, format="PNG")
    return output.getvalue()


SKILLS = [
    "AWS", "GCP", "Azure", "GitHub Actions", "Infrastructure as Code",
    "CI/CD", "CloudFormation",
    "Kubernetes", "Terraform", "Jenkins", "Docker", "Ansible", "Helm",
    "Prometheus", "Grafana", "Datadog", "CloudWatch", "OpenTelemetry",
    "Python", "Java", "JavaScript", "TypeScript", "C#", "C++", "Go",
    "Ruby", "PHP", "Kotlin", "Swift", "Scala", "Dart", "R", "Shell", "Bash",
    "PowerShell", "HTML", "CSS", "Sass", "Bootstrap", "Tailwind CSS",
    "React", "Angular", "Vue", "Next.js", "Node.js", "Express.js", "Redux",
    "Django", "Flask", "FastAPI", "Spring Boot", "Spring", "Hibernate",
    "JPA", "JDBC", "Servlets", "JSP", ".NET", "ASP.NET", "REST API",
    "GraphQL", "WebSocket", "JWT", "OAuth", "Microservices",
    "Flutter", "React Native", "Android", "iOS",
    "Linux", "Ubuntu", "Windows", "Git", "GitHub", "GitLab",
    "Bitbucket", "MySQL", "PostgreSQL", "MongoDB", "Redis", "Oracle",
    "DynamoDB", "SQLite", "SQL Server", "SQL", "NoSQL", "Kafka", "RabbitMQ",
    "Nginx", "Apache",
    "Tomcat", "EC2", "IAM", "S3", "RDS", "Lambda", "Route53", "SQS", "SNS",
    "KMS", "VPC", "VPN", "CDN",
    "Agile", "Scrum", "Jira", "Selenium", "Pytest", "JUnit", "SonarQube",
    "Cypress", "Playwright", "Postman", "Maven", "Gradle", "npm", "Yarn",
    "Machine Learning", "Artificial Intelligence", "Generative AI", "LLM",
    "RAG", "Data Science", "Data Engineering", "Pandas", "NumPy", "SciPy",
    "Scikit-learn", "TensorFlow", "PyTorch", "Keras", "Spark", "Hadoop",
    "Airflow", "Databricks", "Power BI", "Tableau", "Excel",
    "Cybersecurity", "DevSecOps", "OWASP", "UI/UX", "Figma",
]

SECTION_HEADINGS = {
    "SUMMARY", "PROFILE", "PROFESSIONAL SUMMARY", "OBJECTIVE", "SKILLS",
    "TECHNICAL SKILLS", "CORE SKILLS", "EXPERIENCE", "WORK EXPERIENCE",
    "PROFESSIONAL EXPERIENCE", "PROJECTS", "EDUCATION", "CERTIFICATIONS",
    "CERTIFICATES", "ACHIEVEMENTS", "LANGUAGES", "INTERNSHIP", "INTERNSHIPS",
    "CORE COMPETENCIES", "KEY COMPETENCIES", "AREAS OF EXPERTISE",
    "JD-MATCHED SKILLS", "ADDITIONAL DETAILS", "QUALIFICATION",
    "QUALIFICATIONS", "ACADEMIC QUALIFICATION", "ACADEMIC QUALIFICATIONS",
}

# Canonical structure copied from profile/Satyam_Dev_Resume_ATS.pdf so every
# generated CV follows the same reference layout.
HEADING_ALIASES = {
    "SUMMARY": "SUMMARY",
    "PROFILE": "SUMMARY",
    "OBJECTIVE": "SUMMARY",
    "PROFESSIONAL SUMMARY": "SUMMARY",
    "SKILLS": "SKILLS",
    "TECHNICAL SKILLS": "SKILLS",
    "CORE SKILLS": "SKILLS",
    "CORE COMPETENCIES": "CORE COMPETENCIES",
    "KEY COMPETENCIES": "CORE COMPETENCIES",
    "AREAS OF EXPERTISE": "CORE COMPETENCIES",
    "JD-MATCHED SKILLS": "JD-MATCHED SKILLS",
    "EXPERIENCE": "PROFESSIONAL EXPERIENCE",
    "WORK EXPERIENCE": "PROFESSIONAL EXPERIENCE",
    "PROFESSIONAL EXPERIENCE": "PROFESSIONAL EXPERIENCE",
    "INTERNSHIP": "INTERNSHIPS",
    "INTERNSHIPS": "INTERNSHIPS",
    "PROJECTS": "PROJECTS",
    "EDUCATION": "EDUCATION",
    "QUALIFICATION": "QUALIFICATION",
    "QUALIFICATIONS": "QUALIFICATION",
    "ACADEMIC QUALIFICATION": "QUALIFICATION",
    "ACADEMIC QUALIFICATIONS": "QUALIFICATION",
    "CERTIFICATIONS": "CERTIFICATIONS",
    "CERTIFICATES": "CERTIFICATIONS",
    "ACHIEVEMENTS": "ACHIEVEMENTS",
    "LANGUAGES": "LANGUAGES",
    "ADDITIONAL DETAILS": "ADDITIONAL DETAILS",
}

REFERENCE_ORDER = [
    "SUMMARY",
    "SKILLS",
    "JD-MATCHED SKILLS",
    "CORE COMPETENCIES",
    "PROFESSIONAL EXPERIENCE",
    "INTERNSHIPS",
    "PROJECTS",
    "EDUCATION",
    "QUALIFICATION",
    "CERTIFICATIONS",
    "ACHIEVEMENTS",
    "LANGUAGES",
    "ADDITIONAL DETAILS",
]

# ATS parsers gain nothing from these blocks in the source CV.
DROP_HEADINGS = {"DECLARATION", "REFERENCES", "DECLARATION:"}

# Sections whose plain lines are prose, not company/degree entry lines.
PROSE_SECTIONS = {"SUMMARY", "SKILLS", "JD-MATCHED SKILLS", "LANGUAGES"}

# Sections built from entries (company/degree line + description lines).
ENTRY_SECTIONS = {
    "PROFESSIONAL EXPERIENCE", "INTERNSHIPS", "PROJECTS", "EDUCATION",
    "QUALIFICATION", "CERTIFICATIONS",
}

# Sections that read best as a plain bullet list.
BULLET_SECTIONS = {"ACHIEVEMENTS", "CORE COMPETENCIES", "ADDITIONAL DETAILS"}

# Skill buckets used to print a grouped, recruiter-friendly SKILLS block.
SKILL_GROUPS: list[tuple[str, tuple[str, ...]]] = [
    ("Cloud Platforms", (
        "aws", "amazon web services", "azure", "gcp", "google cloud", "ec2",
        "s3", "iam", "rds", "lambda", "route53", "route 53", "sqs", "sns",
        "kms", "vpc", "cloudformation", "cloud deployment", "cloud",
    )),
    ("DevOps & Automation", (
        "ci/cd", "cicd", "infrastructure as code", "iac", "terraform",
        "jenkins", "docker", "kubernetes", "ansible", "helm",
        "github actions", "gitlab ci", "configuration management",
        "infrastructure automation", "automation", "deployment", "devsecops",
    )),
    ("Monitoring & Observability", (
        "prometheus", "grafana", "datadog", "cloudwatch", "opentelemetry",
        "monitoring", "logging",
    )),
    ("Programming & Scripting", (
        "java", "python", "javascript", "typescript", "c#", "c++", "go",
        "golang", "kotlin", "swift", "scala", "ruby", "php", "dart", "shell",
        "bash", "powershell",
    )),
    ("Frameworks & Libraries", (
        "spring boot", "spring", "hibernate", "jpa", "jdbc", "servlets",
        "jsp", "django", "flask", "fastapi", "react", "angular", "vue",
        "next.js", "node.js", "express.js", "redux", ".net", "asp.net",
        "rest api", "graphql", "microservices", "html", "css", "bootstrap",
        "tailwind", "flutter", "react native",
    )),
    ("Databases", (
        "mysql", "postgresql", "mongodb", "redis", "oracle", "dynamodb",
        "sqlite", "sql", "nosql", "kafka", "rabbitmq",
    )),
    ("Tools & Platforms", (
        "git", "github", "gitlab", "bitbucket", "linux", "ubuntu", "windows",
        "maven", "gradle", "npm", "yarn", "jira", "postman", "nginx",
        "apache", "tomcat", "figma", "development tools",
    )),
    ("Data & AI", (
        "machine learning", "artificial intelligence", "generative ai", "llm",
        "rag", "data science", "data engineering", "pandas", "numpy",
        "scikit-learn", "tensorflow", "pytorch", "spark", "hadoop", "airflow",
        "power bi", "tableau", "excel",
    )),
    ("Practices", (
        "agile", "scrum", "owasp", "cybersecurity", "ui/ux", "testing",
        "unit testing", "code review",
    )),
]

# Professional phrasing for competency bullets built from the user's own words.
COMPETENCY_TEMPLATES: list[tuple[tuple[str, ...], str]] = [
    (("infrastructure as code", "iac", "terraform", "cloudformation"),
     "Provisioning and managing infrastructure using Infrastructure as Code "
     "(IaC) for repeatable, version-controlled environments."),
    (("ci/cd", "cicd", "jenkins", "github actions", "pipeline"),
     "Building and maintaining CI/CD pipelines that automate build, test, and "
     "deployment stages."),
    (("configuration management", "ansible", "puppet", "chef"),
     "Standardising server and application configuration to keep environments "
     "consistent and reproducible."),
    (("infrastructure automation", "automation", "scripting", "shell", "bash"),
     "Automating repetitive build, deployment, and operational tasks to reduce "
     "manual effort and human error."),
    (("cloud deployment", "deployment process", "deployment", "release"),
     "Planning and executing application deployments with repeatable, "
     "low-downtime release processes."),
    (("aws", "azure", "gcp", "cloud"),
     "Working with cloud platform services for compute, storage, networking, "
     "and access management."),
    (("docker", "container"),
     "Containerising applications for portable builds across development and "
     "production environments."),
    (("kubernetes", "helm", "orchestration"),
     "Deploying and operating containerised workloads on Kubernetes, including "
     "configuration and scaling."),
    (("monitoring", "prometheus", "grafana", "datadog", "cloudwatch"),
     "Setting up monitoring and alerting to track system health and respond to "
     "issues early."),
    (("git", "github", "gitlab", "version control"),
     "Managing source code with Git-based version control, branching, and pull "
     "request reviews."),
    (("linux", "ubuntu", "server administration"),
     "Administering Linux servers, including packages, permissions, services, "
     "and troubleshooting."),
    (("java", "spring", "hibernate", "j2ee"),
     "Developing and maintaining Java applications using object-oriented "
     "design and framework-based development."),
    (("python", "django", "flask", "fastapi"),
     "Developing Python services and automation scripts with clean, "
     "maintainable code."),
    (("react", "angular", "vue", "javascript", "typescript", "frontend"),
     "Building responsive front-end interfaces and integrating them with "
     "backend APIs."),
    (("rest api", "api", "microservices", "backend"),
     "Designing and consuming REST APIs and service integrations."),
    (("sql", "mysql", "postgresql", "mongodb", "database"),
     "Designing database schemas and writing optimised queries for "
     "application data access."),
    (("development tools", "tooling", "infrastructure setup"),
     "Setting up development tooling and project infrastructure to support "
     "smooth team delivery."),
    (("agile", "scrum", "sprint"),
     "Working in Agile/Scrum teams with sprint planning, reviews, and "
     "collaborative delivery."),
    (("testing", "junit", "pytest", "selenium", "qa"),
     "Writing and running automated tests to protect functionality during "
     "changes."),
]

MAX_COMPETENCY_BULLETS = 8

# Words that already read as a job title, so the summary needs no suffix.
ROLE_NOUNS = (
    "developer", "engineer", "analyst", "administrator", "architect",
    "consultant", "manager", "specialist", "designer", "tester", "intern",
    "trainee", "lead", "scientist", "programmer", "sre", "devops", "support",
)


def extract_pdf_text(pdf_bytes: bytes) -> str:
    """Extract selectable text from an uploaded PDF."""
    reader = PdfReader(BytesIO(pdf_bytes))
    pages = [(page.extract_text() or "").strip() for page in reader.pages]
    text = "\n".join(page for page in pages if page)
    text = clean_text(text)
    if len(text) < 80:
        raise ValueError(
            "No readable text was found in this PDF. Please upload a "
            "selectable-text CV instead of a scanned image PDF."
        )
    return text


BULLET_GLYPHS = "\u2022\u2023\u2043\u2219\u25aa\u25cf\u25e6\u00b7\u007f\uf0a7\uf0b7\uf06c"


# A PDF stores one line per rendered row, so a single bullet or sentence is
# usually split across several lines. Anything shorter than this is treated as
# a deliberate short line instead of a wrapped remainder.
WRAP_MIN_LENGTH = 55
SENTENCE_END = (".", ":", ";", "!", "?")


def _is_numbered_item(line: str) -> bool:
    """True for a user-written list item such as '1. Built the pipeline'."""
    return bool(re.match(r"^\(?\d{1,2}[.)]\s+\S", line.strip()))


def _is_list_item(line: str) -> bool:
    return line.startswith(("-", "*", "\u2022")) or _is_numbered_item(line)


def _starts_new_block(line: str) -> bool:
    """True when a line begins a heading, list item, entry title or date."""
    key = line.upper().rstrip(":")
    if key in SECTION_HEADINGS or key in HEADING_ALIASES or key in DROP_HEADINGS:
        return True
    if _is_list_item(line):
        return True
    return _is_date_line(line) or _is_entry_heading(line)


def _accepts_continuation(line: str) -> bool:
    """Long list items and prose can absorb a wrapped remainder; titles cannot."""
    if len(line) < WRAP_MIN_LENGTH:
        return False
    key = line.upper().rstrip(":")
    if key in SECTION_HEADINGS or key in HEADING_ALIASES or key in DROP_HEADINGS:
        return False
    if _is_date_line(line):
        return False
    if _is_list_item(line):
        return True
    return not _is_entry_heading(line)


def _merge_wrapped_lines(lines: list[str]) -> list[str]:
    """Join lines the PDF wrapped so each bullet stays one complete point."""
    merged: list[str] = []
    for line in lines:
        previous = merged[-1] if merged else ""
        if (
            previous
            and line
            and not _starts_new_block(line)
            and _accepts_continuation(previous)
            and (line[:1].islower() or not previous.endswith(SENTENCE_END))
        ):
            merged[-1] = previous + " " + line
            continue
        merged.append(line)
    return merged


def clean_text(text: str) -> str:
    for glyph in BULLET_GLYPHS:
        text = text.replace(glyph, "-")
    text = text.replace("\u2013", "-").replace("\u2014", "-")
    text = text.replace("\xa0", " ")
    lines: list[str] = []
    blank = False
    for raw in text.splitlines():
        line = re.sub(r"[ \t]+", " ", raw).strip()
        if not line:
            if lines and not blank:
                lines.append("")
            blank = True
            continue
        blank = False
        lines.append(line)
    return "\n".join(_merge_wrapped_lines(lines)).strip()


def save_existing_cv(pdf_bytes: bytes, original_name: str, text: str) -> dict:
    DATA_DIR.mkdir(parents=True, exist_ok=True)
    CV_PDF.write_bytes(pdf_bytes)
    CV_TEXT.write_text(text, encoding="utf-8")
    meta = {
        "original_name": original_name,
        "uploaded_at": datetime.now().isoformat(timespec="seconds"),
        "characters": len(text),
    }
    CV_META.write_text(json.dumps(meta, indent=2), encoding="utf-8")
    return meta


def load_existing_cv() -> tuple[dict, str, bytes] | None:
    if not (CV_PDF.exists() and CV_TEXT.exists() and CV_META.exists()):
        return None
    try:
        meta = json.loads(CV_META.read_text(encoding="utf-8"))
        return meta, CV_TEXT.read_text(encoding="utf-8"), CV_PDF.read_bytes()
    except (OSError, ValueError, json.JSONDecodeError):
        return None


def _contains(text: str, term: str) -> bool:
    aliases = {
        "AWS": ["amazon web services", "aws"],
        "Azure": ["microsoft azure", "azure"],
        "GCP": ["google cloud platform", "gcp"],
        "Go": ["golang", "go"],
        "Route53": ["route 53", "route53"],
        "Artificial Intelligence": ["artificial intelligence", "ai"],
        "Machine Learning": ["machine learning", "ml"],
        "Generative AI": ["generative ai", "genai", "gen ai"],
        "LLM": ["llm", "large language model", "large language models"],
        "Kubernetes": ["kubernetes", "k8s"],
        "Infrastructure as Code": ["infrastructure as code", "iac"],
        "Node.js": ["node.js", "nodejs", "node js"],
        "Next.js": ["next.js", "nextjs", "next js"],
        "Express.js": ["express.js", "expressjs", "express js"],
        "JavaScript": ["javascript", "java script", "js"],
        "TypeScript": ["typescript", "type script", "ts"],
        "REST API": ["rest api", "restful api", "restful services"],
        "Spring Boot": ["spring boot", "springboot"],
        "PostgreSQL": ["postgresql", "postgres"],
        "Scikit-learn": ["scikit-learn", "sklearn"],
        "Power BI": ["power bi", "powerbi"],
        "HTML": ["html", "html5"],
        "CSS": ["css", "css3"],
    }
    options = aliases.get(term, [term])
    return any(
        re.search(r"(?<!\w)" + re.escape(option) + r"(?!\w)", text, re.IGNORECASE)
        for option in options
    )


# Words that fill every JD but never describe a hiring skill.
JD_NOISE_WORDS = {
    "a", "an", "and", "the", "of", "in", "on", "with", "for", "to", "or", "as",
    "at", "by", "from", "is", "are", "be", "will", "must", "should", "can",
    "have", "having", "using", "use", "used", "such", "including", "include",
    "includes", "etc", "we", "you", "your", "our", "their", "this", "that",
    "these", "those", "per", "across", "within", "other", "various", "multiple",
    "new", "experience", "experiences", "knowledge", "understanding", "strong",
    "good", "excellent", "ability", "abilities", "skill", "skills", "proficient",
    "proficiency", "familiarity", "familiar", "expertise", "exposure", "hands",
    "year", "years", "yrs", "minimum", "plus", "preferred", "required",
    "requirement", "requirements", "responsibility", "responsibilities",
    "candidate", "candidates", "role", "roles", "team", "teams", "work",
    "working", "works", "company", "business", "project", "projects",
    "development", "developing", "design", "designing", "support", "supporting",
    "solution", "solutions", "product", "products", "service", "services",
    "client", "clients", "customer", "customers", "stakeholder", "stakeholders",
    "communication", "collaboration", "teamwork", "leadership", "problem",
    "solving", "written", "verbal", "interpersonal", "attention", "detail",
    "fast", "paced", "environment", "degree", "bachelor", "bachelors", "master",
    "masters", "field", "related", "equivalent", "job", "description",
    "position", "opportunity", "apply", "join", "looking", "seeking", "ideal",
    "successful", "applicant", "nice", "must-have", "location", "notice",
    "period", "salary", "ctc", "immediate", "joiner", "shift", "office",
}

# Acronyms that look technical but are just English in caps.
ACRONYM_NOISE = {
    "AND", "OR", "THE", "FOR", "WITH", "YOU", "WE", "OUR", "ALL", "NEW", "JOB",
    "CV", "HR", "JD", "US", "UK", "USA", "EU", "AM", "PM", "OK", "NOT", "ARE",
    "WILL", "CAN", "MUST", "ANY", "ONE", "TWO", "PER", "CTC", "LPA", "WFH",
    "WFO", "NA", "N/A", "MS", "BE", "BS", "MBA", "BCA", "MCA", "BTECH", "MTECH",
}

# Phrases that introduce a list of real skills in a JD.
JD_CUE_RE = re.compile(
    r"(?:experience\s+(?:with|in|on|of)|hands[- ]on\s+(?:experience\s+)?"
    r"(?:with|in)?|knowledge\s+of|proficien(?:t|cy)\s+(?:in|with)|"
    r"expertise\s+(?:in|with)|familiarity\s+with|exposure\s+to|"
    r"working\s+knowledge\s+of|skills?\s*[:\-]|technologies\s*[:\-]|"
    r"tech\s*stack\s*[:\-]|stack\s*[:\-]|tools?\s*[:\-]|must\s*have\s*[:\-]?|"
    r"good\s*to\s*have\s*[:\-]?)",
    re.IGNORECASE,
)

DOTTED_TECH_RE = re.compile(
    r"\b[A-Za-z][A-Za-z0-9]*(?:\.[A-Za-z0-9]+)+\b|\bC\+\+|\bC#|\bF#"
)
ACRONYM_RE = re.compile(r"\b[A-Z][A-Z0-9+#/]{1,7}\b")
MAX_JD_KEYWORDS = 30


def _clean_keyword(phrase: str) -> str:
    """Trim a raw JD fragment down to a usable skill phrase."""
    text = re.sub(r"\([^)]*\)", " ", phrase)
    text = re.sub(r"[^A-Za-z0-9+#./ -]", " ", text)
    text = re.sub(r"\s+", " ", text).strip(" -.")
    if not text:
        return ""
    words = text.split()
    while words and words[0].lower() in JD_NOISE_WORDS:
        words.pop(0)
    while words and words[-1].lower() in JD_NOISE_WORDS:
        words.pop()
    if not words or len(words) > 3 or len(" ".join(words)) > 30:
        return ""
    if all(word.lower() in JD_NOISE_WORDS for word in words):
        return ""
    if not any(char.isalpha() for char in " ".join(words)):
        return ""
    return " ".join(words)


def extract_jd_keywords(jd_text: str) -> list[str]:
    """Pull skill terms out of a JD, including ones outside the built-in list."""
    found: list[str] = []

    for cue in JD_CUE_RE.finditer(jd_text):
        tail = jd_text[cue.end():]
        tail = re.split(r"[.;\n]", tail, maxsplit=1)[0]
        for fragment in re.split(r",|/|\||&|\band\b|\bor\b", tail, flags=re.I):
            keyword = _clean_keyword(fragment)
            if keyword:
                found.append(keyword)

    for token in DOTTED_TECH_RE.findall(jd_text):
        keyword = _clean_keyword(token)
        if keyword:
            found.append(keyword)

    for token in ACRONYM_RE.findall(jd_text):
        if token in ACRONYM_NOISE or token.lower() in JD_NOISE_WORDS:
            continue
        found.append(token)

    ranked = sorted(
        _dedupe(found),
        key=lambda word: -len(re.findall(
            r"(?<!\w)" + re.escape(word) + r"(?!\w)", jd_text, re.IGNORECASE
        )),
    )
    return ranked[:MAX_JD_KEYWORDS]


def match_jd(cv_text: str, jd_text: str, include_dynamic: bool = True) -> dict:
    known = [skill for skill in SKILLS if _contains(jd_text, skill)]
    requested = list(known)
    if include_dynamic:
        seen = {skill.casefold() for skill in known}
        for keyword in extract_jd_keywords(jd_text):
            if keyword.casefold() not in seen:
                seen.add(keyword.casefold())
                requested.append(keyword)
    present = [skill for skill in requested if _contains(cv_text, skill)]
    missing = [skill for skill in requested if skill not in present]
    score = round(100 * len(present) / len(requested)) if requested else 0
    return {
        "requested": _dedupe(requested),
        "matched": _dedupe(present),
        "missing": _dedupe(missing),
        "score": score,
    }


def _dedupe(items: list[str]) -> list[str]:
    seen: set[str] = set()
    result = []
    for item in items:
        key = item.casefold()
        if key not in seen:
            seen.add(key)
            result.append(item)
    return result


ROLE_PHRASE_RE = re.compile(
    r"\b((?:[A-Za-z+#./]+[ -]){0,3}(?:developer|engineer|analyst|administrator"
    r"|architect|consultant|manager|specialist|designer|tester|programmer"
    r"|scientist|intern|trainee))\b",
    re.IGNORECASE,
)
JD_TITLE_LINE_RE = re.compile(
    r"^\s*(?:job\s*title|job\s*role|role|position|designation)\s*[:\-]\s*(.+)$",
    re.IGNORECASE | re.MULTILINE,
)
JD_TITLE_CUE_RE = re.compile(
    r"(?:hiring for|looking for|seeking|position of|role of|we need)\s+"
    r"(?:an?\s+)?([A-Za-z+#./ -]{3,60})",
    re.IGNORECASE,
)
# Seniority words that should not decide whether two titles match.
TITLE_NOISE = {
    "senior", "junior", "sr", "jr", "lead", "principal", "staff", "associate",
    "i", "ii", "iii", "iv", "1", "2", "3", "fresher", "entry", "level",
}


# Words that sit in front of a title in a JD sentence but are not part of it,
# e.g. "We are hiring a DevOps Engineer".
TITLE_FILLER = {
    "a", "an", "the", "we", "are", "is", "for", "our", "of", "as", "to", "and",
    "hiring", "looking", "seeking", "need", "needs", "want", "wanted", "join",
    "urgent", "urgently", "immediate", "immediately", "required", "requires",
    "requirement", "opening", "openings", "vacancy", "job", "role", "position",
    "title", "designation", "experienced", "fresher", "new", "one", "multiple",
}


def _role_phrase(text: str) -> str:
    match = ROLE_PHRASE_RE.search(text)
    if not match:
        return ""
    words = re.sub(r"\s+", " ", match.group(1)).strip().split()
    while words and words[0].strip(".,:-()").lower() in TITLE_FILLER:
        words.pop(0)
    return " ".join(words)


def jd_role_title(jd_text: str) -> str:
    """Best guess at the job title the JD is hiring for."""
    line_match = JD_TITLE_LINE_RE.search(jd_text)
    if line_match:
        candidate = _role_phrase(line_match.group(1)) or line_match.group(1)
        return re.sub(r"\s+", " ", candidate).strip()[:60]
    cue_match = JD_TITLE_CUE_RE.search(jd_text)
    if cue_match:
        candidate = _role_phrase(cue_match.group(1))
        if candidate:
            return candidate
    for line in [row.strip() for row in jd_text.splitlines() if row.strip()][:6]:
        candidate = _role_phrase(line)
        if candidate:
            return candidate
    return ""


def cv_role_title(cv_text: str) -> str:
    """Role headline already present in the CV."""
    return _cv_identity(cv_text)[1]


def _title_tokens(title: str) -> set[str]:
    words = re.findall(r"[A-Za-z+#.]+", title.lower())
    return {word for word in words if word not in TITLE_NOISE}


def title_alignment(cv_text: str, jd_text: str) -> dict:
    """Compare the CV headline with the JD title, an ATS ranking signal."""
    jd_title = jd_role_title(jd_text)
    cv_title = cv_role_title(cv_text)
    jd_tokens = _title_tokens(jd_title)
    cv_tokens = _title_tokens(cv_title)
    if not jd_tokens:
        return {"jd_title": jd_title, "cv_title": cv_title, "score": None}
    if not cv_tokens:
        return {"jd_title": jd_title, "cv_title": cv_title, "score": 0}
    overlap = jd_tokens & cv_tokens
    return {
        "jd_title": jd_title,
        "cv_title": cv_title,
        "score": round(100 * len(overlap) / len(jd_tokens)),
    }


YEAR_REQUIREMENT_RE = re.compile(
    r"(\d{1,2})\s*(?:\+|plus)?\s*(?:[-\u2013]|to)?\s*(?:\d{1,2})?\s*\+?\s*"
    r"(?:years?|yrs?)",
    re.IGNORECASE,
)
MONTH_INDEX = {
    "jan": 1, "feb": 2, "mar": 3, "apr": 4, "may": 5, "jun": 6, "jul": 7,
    "aug": 8, "sep": 9, "oct": 10, "nov": 11, "dec": 12,
}
DATE_POINT = (
    r"(?:(?:jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec)[a-z]*\.?\s*,?\s*)?"
    r"(?:\d{1,2}[/-])?(?:19|20)\d{2}"
)
DATE_RANGE_RE = re.compile(
    rf"({DATE_POINT})\s*(?:[-\u2013\u2014]|to|till|until)\s*"
    rf"(present|current|now|till date|{DATE_POINT})",
    re.IGNORECASE,
)


def jd_required_years(jd_text: str) -> int | None:
    """Minimum years of experience the JD asks for, if it states any."""
    values = [
        int(match.group(1))
        for match in YEAR_REQUIREMENT_RE.finditer(jd_text)
        if 0 < int(match.group(1)) <= 30
    ]
    return max(values) if values else None


def _month_number(token: str) -> int | None:
    token = token.strip().lower()
    if not token:
        return None
    month = 1
    month_word = re.search(r"[a-z]{3}", token)
    if month_word:
        month = MONTH_INDEX.get(month_word.group(0), 1)
    elif re.match(r"^\d{1,2}[/-]", token):
        month = int(re.split(r"[/-]", token)[0])
    year_match = re.search(r"(19|20)\d{2}", token)
    if not year_match:
        return None
    return int(year_match.group(0)) * 12 + max(1, min(12, month))


def cv_experience_years(cv_text: str) -> float:
    """Total professional months in the CV, overlapping roles counted once."""
    _, sections, _ = parse_cv_sections(cv_text)
    text = "\n".join(
        sections.get("PROFESSIONAL EXPERIENCE", [])
        + sections.get("INTERNSHIPS", [])
    )
    now = datetime.now()
    today = now.year * 12 + now.month
    spans: list[tuple[int, int]] = []
    for match in DATE_RANGE_RE.finditer(text):
        start = _month_number(match.group(1))
        end_token = match.group(2).lower()
        end = (
            today
            if end_token in {"present", "current", "now", "till date"}
            else _month_number(end_token)
        )
        if start and end and end >= start:
            spans.append((start, min(end, today)))

    total = 0
    last_end = 0
    for start, end in sorted(spans):
        start = max(start, last_end)
        if end > start:
            total += end - start
            last_end = end
    return round(total / 12, 1)


def ats_report(cv_text: str, jd_text: str, match: dict | None = None) -> dict:
    """Score a CV the way a hiring ATS does: skills, title, tenure, format."""
    match = match or match_jd(cv_text, jd_text)
    title = title_alignment(cv_text, jd_text)
    required_years = jd_required_years(jd_text)
    have_years = cv_experience_years(cv_text)

    _, sections, _ = parse_cv_sections(cv_text)
    format_checks = {
        "Email address": bool(re.search(r"[^@\s]+@[^@\s]+\.[A-Za-z]{2,}", cv_text)),
        "Phone number": bool(re.search(r"(?<!\d)(?:\+?\d[\d ()-]{7,}\d)(?!\d)", cv_text)),
        "Summary section": bool(sections.get("SUMMARY")),
        "Skills section": bool(sections.get("SKILLS")),
        "Experience or projects": bool(
            sections.get("PROFESSIONAL EXPERIENCE")
            or sections.get("INTERNSHIPS")
            or sections.get("PROJECTS")
        ),
        "Education section": bool(
            sections.get("EDUCATION") or sections.get("QUALIFICATION")
        ),
    }
    format_score = round(
        100 * sum(format_checks.values()) / len(format_checks)
    )

    years_score = None
    if required_years:
        years_score = min(100, round(100 * have_years / required_years))

    parts = [(match["score"], 45), (format_score, 20)]
    if title["score"] is not None:
        parts.append((title["score"], 20))
    if years_score is not None:
        parts.append((years_score, 15))
    weight = sum(share for _, share in parts)
    overall = round(sum(value * share for value, share in parts) / weight)

    suggestions: list[str] = []
    if match["missing"]:
        suggestions.append(
            "The JD asks for "
            + ", ".join(match["missing"][:8])
            + ". Tick only the ones you genuinely have so they can be added."
        )
    if title["score"] is not None and title["score"] < 60 and title["jd_title"]:
        suggestions.append(
            f"The JD hires a \"{title['jd_title']}\" but your CV headline reads "
            f"\"{title['cv_title'] or 'no clear role'}\". Match the headline if "
            "the role really is the same."
        )
    if required_years and have_years + 0.5 < required_years:
        suggestions.append(
            f"The JD wants about {required_years} years and your CV dates add up "
            f"to {have_years}. Write clear month-year ranges for every role."
        )
    for label, ok in format_checks.items():
        if not ok:
            suggestions.append(f"{label} is missing — ATS parsers look for it.")

    return {
        "overall": overall,
        "skills_score": match["score"],
        "title": title,
        "required_years": required_years,
        "cv_years": have_years,
        "years_score": years_score,
        "format_score": format_score,
        "format_checks": format_checks,
        "suggestions": suggestions,
    }


def _prioritise_entry_bullets(lines: list[str], keywords: list[str]) -> list[str]:
    """Within each role, list JD-relevant bullets first; keep roles in order."""
    if not keywords:
        return lines
    groups: list[list[str]] = []
    for line in lines:
        if _is_list_item(line) and groups:
            groups[-1].append(line)
        else:
            groups.append([line])

    ordered: list[str] = []
    for group in groups:
        head, bullets = group[0], group[1:]
        ordered.append(head)
        if any(_is_numbered_item(bullet) for bullet in bullets):
            # A numbered list states its own sequence, so leave it alone.
            ordered.extend(bullets)
            continue
        ordered.extend(
            sorted(
                bullets,
                key=lambda bullet: -sum(
                    _contains(bullet, keyword) for keyword in keywords
                ),
            )
        )
    return ordered


def build_tailored_text(
    cv_text: str,
    match: dict,
    extra_skills: list[str] | None = None,
    headline: str | None = None,
) -> str:
    """Prioritise verified JD skills without adding unsupported claims."""
    header, sections, _ = parse_cv_sections(cv_text)
    confirmed = [skill for skill in (extra_skills or []) if skill.strip()]
    matched = _dedupe(list(match["matched"]) + confirmed)

    if confirmed:
        # Match the reference CV's "- Category: items" skill lines instead of
        # dropping a bare comma list into the section.
        sections.setdefault("SKILLS", [])
        sections["SKILLS"].append("- JD Skills: " + ", ".join(confirmed))

    if matched:
        # Keep every original skill line, but move JD-relevant lines first.
        sections["SKILLS"] = sorted(
            sections.get("SKILLS", []),
            key=lambda line: -sum(_contains(line, skill) for skill in matched),
        )
    # Matched skills already appear in SKILLS, so a second keyword-only block
    # would just repeat them and stop the CV looking like the reference layout.
    sections.pop("JD-MATCHED SKILLS", None)

    for name in ("PROFESSIONAL EXPERIENCE", "INTERNSHIPS", "PROJECTS"):
        if sections.get(name):
            sections[name] = _prioritise_entry_bullets(
                _normalize_section_lines(name, sections[name]), matched
            )

    if headline and headline.strip():
        # The reference CV prints the role in capitals right under the name.
        cleaned = re.sub(r"\s+", " ", headline).strip().upper()
        if len(header) >= 2:
            header[1] = cleaned
        else:
            header.append(cleaned)

    rebuilt = list(header)
    for heading in REFERENCE_ORDER:
        lines = sections.get(heading, [])
        if lines:
            rebuilt.extend(("", heading, *lines))

    return restructure_to_reference("\n".join(rebuilt))


def parse_cv_sections(text: str) -> tuple[list[str], dict[str, list[str]], list[str]]:
    """Split any CV text into header lines and canonical sections."""
    lines = clean_text(text).splitlines()
    header: list[str] = []
    sections: dict[str, list[str]] = {}
    order: list[str] = []
    current: str | None = None
    dropping = False

    for raw in lines:
        line = raw.strip()
        if not line:
            continue
        key = line.upper().rstrip(":")
        if key in DROP_HEADINGS:
            dropping = True
            current = None
            continue
        canonical = HEADING_ALIASES.get(key)
        if canonical:
            dropping = False
            current = canonical
            if canonical not in sections:
                sections[canonical] = []
                order.append(canonical)
            continue
        if dropping:
            continue
        if current is None:
            header.append(line)
        else:
            sections[current].append(line)

    return header, sections, order


MONTH_WORDS = (
    "january", "february", "march", "april", "may", "june", "july", "august",
    "september", "october", "november", "december", "jan", "feb", "mar", "apr",
    "jun", "jul", "aug", "sep", "sept", "oct", "nov", "dec",
)


def _is_date_line(line: str) -> bool:
    """True for a stand-alone duration line such as 'July 2024 - March 2025'."""
    text = line.strip().rstrip(".")
    if not text or text.startswith(("-", "*", "\u2022")) or _is_numbered_item(text):
        return False
    words = text.split()
    if len(words) > 8:
        return False
    low = text.lower()
    has_period = bool(re.search(r"(19|20)\d{2}", text)) or "present" in low
    if not has_period:
        return False
    allowed = set(MONTH_WORDS) | {"present", "current", "to", "till", "-", "–"}
    for word in words:
        token = word.strip("(),.-–|").lower()
        if not token:
            continue
        if token in allowed or token.isdigit() or re.fullmatch(r"\d{1,2}/\d{2,4}", token):
            continue
        return False
    return True


# Sections where the reference CV highlights only the degree or course name and
# leaves the institute and the years in normal weight.
TITLE_DETAIL_SECTIONS = {"EDUCATION", "QUALIFICATION", "CERTIFICATIONS"}


def _split_entry_detail(line: str, section: str) -> tuple[str, str]:
    """Split 'Degree - Institute (years)' into highlighted and plain parts."""
    if section not in TITLE_DETAIL_SECTIONS:
        return line, ""
    parts = re.split(r"\s[-\u2013\u2014]\s", line, maxsplit=1)
    if len(parts) != 2 or not parts[1].strip():
        return line, ""
    return parts[0].strip(), parts[1].strip()


def _is_entry_heading(line: str) -> bool:
    """True for a company / degree / project title line, not a description."""
    text = line.strip()
    if not text or text.startswith(("-", "*", "\u2022")) or _is_numbered_item(text):
        return False
    words = text.split()
    if len(words) > 16:
        return False
    if re.search(r"(19|20)\d{2}", text):
        return True
    separated = bool(re.search(r"\s[-\u2013\u2014|]\s", text))
    if separated and "present" in text.lower():
        return True
    if separated and len(words) <= 14 and not text.endswith("."):
        return True
    return len(words) <= 6 and text == text.upper() and not text.endswith(".")


def _normalize_section_lines(name: str, lines: list[str]) -> list[str]:
    """Keep entry titles on their own line and turn descriptions into bullets."""
    if name in PROSE_SECTIONS:
        return lines
    rows: list[str] = []
    if name in BULLET_SECTIONS:
        for line in lines:
            if _is_numbered_item(line):
                rows.append(line)
                continue
            body = line.lstrip("-*\u2022 ").strip()
            if body:
                rows.append("- " + body)
        return rows
    if name not in ENTRY_SECTIONS:
        return lines
    entry_seen = False
    for line in lines:
        if _is_numbered_item(line):
            entry_seen = True
            rows.append(line)
            continue
        if line.startswith(("-", "*", "\u2022")):
            body = line.lstrip("-*\u2022 ").strip()
            if body:
                rows.append("- " + body)
            continue
        if _is_entry_heading(line):
            entry_seen = True
            rows.append(line)
        elif entry_seen and len(line.split()) >= 4:
            rows.append("- " + line)
        else:
            entry_seen = True
            rows.append(line)
    return rows


def restructure_to_reference(text: str) -> str:
    """Reorder any CV text into the reference resume layout, ATS-safe."""
    header, sections, order = parse_cv_sections(text)

    out: list[str] = header[:3] if header else []
    leftover_header = header[3:]
    if leftover_header:
        sections.setdefault("ADDITIONAL DETAILS", [])
        if "ADDITIONAL DETAILS" not in order:
            order.append("ADDITIONAL DETAILS")
        sections["ADDITIONAL DETAILS"] = leftover_header + sections["ADDITIONAL DETAILS"]

    ordered = [name for name in REFERENCE_ORDER if sections.get(name)]
    ordered += [name for name in order if name not in ordered and sections.get(name)]

    for name in ordered:
        out += ["", name]
        out += _normalize_section_lines(name, sections[name])

    return "\n".join(out).strip() + "\n"


def detect_skills(text: str) -> list[str]:
    """Return known skills that actually appear in the given text."""
    return _dedupe([skill for skill in SKILLS if _contains(text, skill)])


ROLE_WORDS = (
    "developer", "engineer", "analyst", "administrator", "architect",
    "consultant", "manager", "specialist", "designer", "tester", "trainee",
    "intern", "lead", "scientist", "programmer", "devops", "support",
)


def _cv_identity(text: str) -> tuple[str, str, str, str]:
    """Extract only identity fields that are visibly present in the CV."""
    header, sections, _ = parse_cv_sections(text)
    all_lines = [line.strip() for line in clean_text(text).splitlines() if line.strip()]
    candidates = header or all_lines[:8]

    email_match = re.search(
        r"\b[A-Z0-9._%+-]+@[A-Z0-9.-]+\.[A-Z]{2,}\b", text, re.IGNORECASE
    )
    phone_match = re.search(
        r"(?<!\d)(?:\+?\d[\d ()-]{7,}\d)(?!\d)", text
    )
    email = email_match.group(0) if email_match else ""
    phone = re.sub(r"\s+", " ", phone_match.group(0)).strip() if phone_match else ""

    def is_contact(line: str) -> bool:
        low = line.lower()
        return (
            "@" in line
            or bool(re.search(r"\d{7,}", re.sub(r"\D", "", line)))
            or "linkedin" in low
            or "github" in low
            or "|" in line
        )

    name = next(
        (
            line
            for line in candidates
            if not is_contact(line)
            and line.upper().rstrip(":") not in SECTION_HEADINGS
            and line.upper() not in {"RESUME", "CURRICULUM VITAE", "CV"}
            and 1 <= len(line.split()) <= 6
        ),
        "Candidate",
    )

    role = next(
        (
            line
            for line in candidates
            if line != name
            and not is_contact(line)
            and any(word in line.lower() for word in ROLE_WORDS)
            and len(line.split()) <= 10
        ),
        "",
    )
    if not role:
        summary = " ".join(sections.get("SUMMARY", [])[:3])
        title_match = re.search(
            r"\b([A-Za-z+#.]+(?:\s+[A-Za-z+#.]+){0,4}\s+"
            r"(?:Developer|Engineer|Analyst|Administrator|Architect|"
            r"Consultant|Specialist|Designer|Tester|Manager))\b",
            summary,
            re.IGNORECASE,
        )
        role = title_match.group(1) if title_match else ""

    return name.title() if name.isupper() else name, role, email, phone


def build_cold_email(
    cv_text: str,
    target_role: str = "",
    company: str = "",
) -> dict[str, str | list[str]]:
    """Create a concise, truthful HR email using only facts found in the CV."""
    header, sections, _ = parse_cv_sections(cv_text)
    del header
    name, inferred_role, email, phone = _cv_identity(cv_text)
    role = target_role.strip() or inferred_role.strip()
    if role and (role.isupper() or role.islower()):
        role = role.title()
    role = re.sub(r"\bDevops\b", "DevOps", role)
    role_label = role or "relevant"
    company_name = company.strip()

    experience_lines = [
        line.strip()
        for line in sections.get("PROFESSIONAL EXPERIENCE", [])
        if line.strip()
    ]
    internship_lines = [
        line.strip()
        for line in sections.get("INTERNSHIPS", [])
        if line.strip()
    ]
    has_experience = bool(experience_lines)
    has_internship = bool(internship_lines)

    years_match = re.search(
        r"\b(\d+(?:\.\d+)?\+?)\s*(?:years?|yrs?)\b",
        " ".join(sections.get("SUMMARY", [])),
        re.IGNORECASE,
    )
    years = years_match.group(1) if years_match else ""
    skills = detect_skills(cv_text)[:6]
    skill_phrase = ", ".join(skills)

    if has_experience:
        status = "Experienced professional"
        experience_phrase = (
            f"I have {years} years of professional experience"
            if years
            else "I have professional experience"
        )
        if role:
            experience_phrase += f" as a {role}"
        experience_phrase += "."
    elif has_internship:
        status = "Fresher with internship experience"
        experience_phrase = (
            f"I am a fresher pursuing {role_label} opportunities, with practical "
            "internship exposure documented in my CV."
        )
    else:
        status = "Fresher"
        experience_phrase = (
            f"I am a fresher seeking {role_label} opportunities and looking to "
            "apply the skills and project knowledge documented in my CV."
        )

    if skills:
        background = f"My core skills include {skill_phrase}."
    else:
        background = (
            "My attached CV outlines my relevant education, projects, and skills."
        )

    destination = f" at {company_name}" if company_name else " in your organization"
    subject_role = role if role else "Relevant"
    subject = f"Application for {subject_role} Opportunities"
    if name != "Candidate":
        subject += f" | {name}"

    signoff = [name]
    if phone:
        signoff.append(phone)
    if email and email.casefold() != name.casefold():
        signoff.append(email)

    body = "\n".join(
        [
            "Dear Hiring Manager,",
            "",
            "I hope you are doing well.",
            "",
            f"My name is {name}. {experience_phrase} {background}",
            "",
            f"I am reaching out to explore suitable {role_label} opportunities"
            f"{destination}. I have attached my CV for your review and would "
            "appreciate the opportunity to discuss how my background could "
            "contribute to your team.",
            "",
            "Thank you for your time and consideration.",
            "",
            "Best regards,",
            *signoff,
        ]
    )

    return {
        "subject": subject,
        "body": body,
        "status": status,
        "name": name,
        "role": role,
        "skills": skills,
    }


def _section_lines(heading: str, content: str) -> list[str]:
    """Keep user's own bullets as bullets; plain lines stay entry/detail lines."""
    text = clean_text(content)
    if not text:
        return []
    rows = ["", heading]
    for raw in text.splitlines():
        line = raw.strip()
        if not line:
            continue
        if line.startswith(("-", "*", "•")):
            rows.append("- " + line.lstrip("-*• ").strip())
        elif line.upper().rstrip(":") in SECTION_HEADINGS:
            rows += ["", line.upper().rstrip(":")]
        else:
            rows.append(line)
    return rows


def _looks_like_skill_list(lines: list[str]) -> bool:
    """True when the free-text lines read like a skill list, not sentences."""
    body = [line for line in lines if line]
    if not body:
        return False
    if len(body) == 1:
        only = body[0].rstrip()
        return not only.endswith(".") and ("," in only or len(only.split()) <= 6)
    short = [
        line
        for line in body
        if len(line.split()) <= 6 and not line.rstrip().endswith(".")
    ]
    return len(short) >= max(2, int(len(body) * 0.7))


def group_skills(items: list[str]) -> list[str]:
    """Print skills as labelled bullet groups, matching the reference CV."""
    if len(items) < 4:
        return ["- " + ", ".join(items)] if items else []
    buckets: dict[str, list[str]] = {}
    leftover: list[str] = []
    for item in items:
        low = item.lower()
        for label, keywords in SKILL_GROUPS:
            if any(word in low for word in keywords):
                buckets.setdefault(label, []).append(item)
                break
        else:
            leftover.append(item)
    rows = [
        f"- {label}: {', '.join(buckets[label])}"
        for label, _ in SKILL_GROUPS
        if buckets.get(label)
    ]
    if leftover:
        rows.append(f"- Additional Skills: {', '.join(leftover)}")
    return rows


def build_competencies(phrases: list[str], skills: list[str]) -> list[str]:
    """Turn the user's own skill phrases into professional competency bullets."""
    source = [phrase.strip() for phrase in phrases + skills if phrase.strip()]
    if not source:
        return []
    bullets: list[str] = []
    for keywords, sentence in COMPETENCY_TEMPLATES:
        if sentence in bullets:
            continue
        if any(word in phrase.lower() for phrase in source for word in keywords):
            bullets.append(sentence)
        if len(bullets) >= MAX_COMPETENCY_BULLETS:
            break
    return bullets


def build_summary(headline: str, skills: list[str], has_history: bool) -> list[str]:
    """Write a professional summary when the user leaves that box empty."""
    role = headline.strip()
    if role and role == role.lower():
        role = role.title()
    if role and not any(noun in role.lower() for noun in ROLE_NOUNS):
        role = f"{role} professional"
    role = role or "Technology professional"
    top = [skill for skill in skills if len(skill) < 40][:5]
    sentences = []
    if has_history:
        sentences.append(
            f"{role} with hands-on project and workplace experience delivering "
            "reliable, well-documented solutions."
        )
    else:
        sentences.append(
            f"{role} with practical, hands-on exposure to modern development "
            "and deployment workflows."
        )
    if top:
        sentences.append(f"Core strengths across {', '.join(top)}.")
    sentences.append(
        "Comfortable owning work end to end — from requirement and design "
        "through implementation, testing, and deployment — and improving "
        "delivery through automation and clear documentation."
    )
    sentences.append(
        "Quick to pick up new tools and keen to contribute in a collaborative, "
        "delivery-focused team."
    )
    return sentences


def compose_cv_text(
    name: str = "",
    email: str = "",
    phone: str = "",
    address: str = "",
    city: str = "",
    state: str = "",
    details: str = "",
    sections: dict[str, str] | None = None,
    headline: str = "",
) -> str:
    """Build an ATS-safe CV from whatever the user filled. No field is required."""
    location = ", ".join(part.strip() for part in (city, state) if part.strip())
    contact_parts = [phone, email, address, location]
    contact = " | ".join(part.strip() for part in contact_parts if part.strip())

    lines = [(name.strip() or "CURRICULUM VITAE").upper()]
    if headline.strip():
        lines.append(headline.strip().upper())
    if contact:
        lines.append(contact)

    filled = sections or {}
    extra = details.strip()

    extra_headed: list[str] = []
    extra_loose: list[str] = []
    current_extra = ""
    for raw in clean_text(extra).splitlines():
        line = raw.strip()
        if not line:
            continue
        upper = line.upper().rstrip(":")
        if upper in SECTION_HEADINGS:
            current_extra = upper
            extra_headed += ["", upper]
            continue
        if line.startswith(("-", "*", "•")):
            line = "- " + line.lstrip("-*• ").strip()
        if current_extra:
            extra_headed.append(line)
        else:
            extra_loose.append(line)

    loose_as_skills = _looks_like_skill_list(
        [line for line in extra_loose if not line.startswith("- ")]
    )

    skill_items: list[str] = []
    if filled.get("skills"):
        skill_items += [
            part.strip()
            for part in re.split(r"[\n,;]+", filled["skills"])
            if part.strip()
        ]
    if loose_as_skills:
        for line in extra_loose:
            skill_items += [
                part.strip()
                for part in re.split(r"[,;]|\s+/\s+", line.lstrip("- "))
                if part.strip()
            ]
    known = detect_skills("\n".join([extra, *filled.values()]))
    joined = " ".join(skill_items).lower()
    skill_items += [skill for skill in known if skill.lower() not in joined]
    skill_items = _dedupe(skill_items)

    has_history = any(
        filled.get(key, "").strip()
        for key in ("experience", "internships", "projects")
    )
    if filled.get("summary"):
        lines += ["", "SUMMARY"]
        lines += [
            line.strip()
            for line in clean_text(filled["summary"]).splitlines()
            if line.strip()
        ]
    elif skill_items or headline.strip():
        lines += ["", "SUMMARY"]
        lines += build_summary(headline, skill_items, has_history)

    if skill_items:
        lines += ["", "SKILLS"]
        lines += group_skills(skill_items)

    competency_source = extra_loose if loose_as_skills else []
    competencies = build_competencies(competency_source, skill_items)
    if competencies and not has_history:
        lines += ["", "CORE COMPETENCIES"]
        lines += ["- " + bullet for bullet in competencies]

    for heading, key in (
        ("PROFESSIONAL EXPERIENCE", "experience"),
        ("PROJECTS", "projects"),
        ("EDUCATION", "education"),
        ("QUALIFICATION", "qualification"),
        ("CERTIFICATIONS", "certifications"),
        ("INTERNSHIPS", "internships"),
        ("ACHIEVEMENTS", "achievements"),
        ("LANGUAGES", "languages"),
    ):
        lines += _section_lines(heading, filled.get(key, ""))

    lines += extra_headed
    if extra_loose and not loose_as_skills:
        lines += ["", "ADDITIONAL DETAILS"]
        lines += [
            line if line.startswith("- ") else "- " + line for line in extra_loose
        ]

    return restructure_to_reference("\n".join(lines))


def prepare_profile_photo(photo_bytes: bytes) -> BytesIO:
    """Crop a portrait photo to a consistent size for the CV header."""
    try:
        image = PILImage.open(BytesIO(photo_bytes))
        image = image.convert("RGB")
    except Exception as exc:
        raise ValueError(
            "Could not read the profile photo. Please upload a JPG, PNG, or WEBP image."
        ) from exc

    target_w, target_h = 340, 420
    src_w, src_h = image.size
    if src_w < 40 or src_h < 40:
        raise ValueError("Profile photo is too small. Please upload a clearer picture.")

    scale = max(target_w / src_w, target_h / src_h)
    resized = image.resize(
        (max(1, int(src_w * scale)), max(1, int(src_h * scale))),
        PILImage.Resampling.LANCZOS,
    )
    left = max(0, (resized.width - target_w) // 2)
    top = max(0, (resized.height - target_h) // 2)
    cropped = resized.crop((left, top, left + target_w, top + target_h))
    output = BytesIO()
    cropped.save(output, format="JPEG", quality=88, optimize=True)
    output.seek(0)
    return output


def _split_header_lines(text: str) -> tuple[list[str], list[str]]:
    """Take the name / role / contact block off the top of the CV."""
    header: list[str] = []
    rest: list[str] = []
    in_body = False
    for raw in text.splitlines():
        line = raw.strip()
        if not in_body:
            if not line:
                continue
            upper = line.upper().rstrip(":")
            if upper in SECTION_HEADINGS or upper == "JD-MATCHED SKILLS" or len(header) >= 3:
                in_body = True
                rest.append(raw)
                continue
            header.append(line)
            continue
        rest.append(raw)
    return header, rest


def random_output_stem() -> str:
    """File name for downloads: My_Agent + 5 random digits."""
    return f"My_Agent_{secrets.randbelow(90000) + 10000}"


def manual_output_stem(name: str) -> str:
    base = re.sub(r"[^A-Za-z0-9]+", "_", name).strip("_") or "My"
    return f"{base}_ATS_CV"


def build_ats_pdf(
    text: str,
    title: str = "ATS Resume",
    photo_bytes: bytes | None = None,
    template: str = "reference",
) -> bytes:
    """Create an attractive, selectable, one-column ATS-safe PDF."""
    output = BytesIO()
    styles = getSampleStyleSheet()
    design = cv_template(template)
    primary = colors.HexColor("#" + design["primary"])
    accent = colors.HexColor("#" + design["accent"])
    heading_fill = colors.HexColor("#" + design["heading_fill"])
    heading_text = colors.HexColor("#" + design["heading_text"])
    ruled = design.get("ruled_headings", False)
    compact = design.get("compact", False)
    body_colour = colors.black if ruled else colors.HexColor("#2A2E3A")
    entry_colour = colors.black if ruled else colors.HexColor("#1F2437")
    gap = 1.5 if compact else 3
    name_style = ParagraphStyle(
        "Name", parent=styles["Normal"], fontName="Helvetica-Bold",
        fontSize=17, leading=19, spaceAfter=2 if compact else 4,
        textColor=primary,
    )
    heading_style = ParagraphStyle(
        "Heading", parent=styles["Normal"], fontName="Helvetica-Bold",
        fontSize=11, leading=13,
        spaceBefore=7 if compact else 11, spaceAfter=1 if compact else 6,
        textColor=heading_text,
        backColor=None if ruled else heading_fill,
        borderPadding=None if ruled else (4, 6, 4, 6),
        borderRadius=None if ruled else 2,
        keepWithNext=1,
    )
    body_style = ParagraphStyle(
        "Body", parent=styles["Normal"], fontName="Helvetica",
        fontSize=9.5, leading=11.4 if compact else 12.2,
        spaceAfter=1 if compact else 2, textColor=body_colour,
    )
    bullet_style = ParagraphStyle(
        "Bullet", parent=body_style, leftIndent=13, firstLineIndent=-9,
        bulletIndent=1, spaceAfter=0.5 if compact else 2,
    )
    role_style = ParagraphStyle(
        "Role", parent=styles["Normal"], fontName="Helvetica-Bold",
        fontSize=11, leading=13, spaceAfter=1 if compact else 3,
        textColor=accent,
    )
    entry_style = ParagraphStyle(
        "Entry", parent=styles["Normal"], fontName="Helvetica-Bold",
        fontSize=10, leading=12,
        spaceBefore=3 if compact else 4, spaceAfter=0.5 if compact else 1,
        textColor=entry_colour, keepWithNext=1,
    )
    date_style = ParagraphStyle(
        "Date", parent=styles["Normal"], fontName="Helvetica-Oblique",
        fontSize=9, leading=11, spaceAfter=1, textColor=body_colour,
        keepWithNext=1,
    )
    entry_detail_style = ParagraphStyle(
        "EntryDetail", parent=entry_style, fontName="Helvetica",
    )

    story = []
    source_lines = text.splitlines()
    nonempty_seen = 0
    if photo_bytes:
        header, rest = _split_header_lines(text)
        left = []
        for index, line in enumerate(header):
            safe = html.escape(line)
            if index == 0:
                left.append(Paragraph(safe, name_style))
            elif index == 1 and line == line.upper() and len(line.split()) <= 6:
                left.append(Paragraph(safe, role_style))
            else:
                left.append(Paragraph(safe, body_style))
        if not left:
            left = [Paragraph(" ", body_style)]
        photo = RLImage(
            prepare_profile_photo(photo_bytes),
            width=1.12 * inch,
            height=1.38 * inch,
        )
        table = Table([[left, photo]], colWidths=[5.78 * inch, 1.22 * inch])
        table.setStyle(
            TableStyle(
                [
                    ("VALIGN", (0, 0), (-1, -1), "TOP"),
                    ("ALIGN", (1, 0), (1, 0), "RIGHT"),
                    ("LEFTPADDING", (0, 0), (-1, -1), 0),
                    ("RIGHTPADDING", (0, 0), (-1, -1), 0),
                    ("TOPPADDING", (0, 0), (-1, -1), 0),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 2),
                ]
            )
        )
        story.append(table)
        story.append(Spacer(1, 4 if compact else 6))
        source_lines = rest
        nonempty_seen = 3

    section = ""
    for raw in source_lines:
        line = raw.strip()
        if not line:
            story.append(Spacer(1, gap))
            continue
        safe = html.escape(line)
        upper = line.upper().rstrip(":")
        if nonempty_seen == 0:
            story.append(Paragraph(safe, name_style))
        elif nonempty_seen == 1 and line == upper and len(line.split()) <= 6:
            story.append(Paragraph(safe, role_style))
        elif upper in SECTION_HEADINGS or upper == "JD-MATCHED SKILLS":
            if ruled and not section:
                story.append(
                    HRFlowable(
                        width="100%", thickness=0.7, color=colors.black,
                        spaceBefore=2, spaceAfter=0,
                    )
                )
            section = HEADING_ALIASES.get(upper, upper)
            story.append(Paragraph(html.escape(upper), heading_style))
            if ruled:
                story.append(
                    HRFlowable(
                        width="100%", thickness=0.7, color=colors.black,
                        spaceBefore=0, spaceAfter=3,
                    )
                )
        elif line.startswith(("-", "*", "•")):
            content = html.escape(line.lstrip("-*• ").strip())
            story.append(Paragraph(f"• {content}", bullet_style))
        elif _is_numbered_item(line):
            story.append(Paragraph(safe, bullet_style))
        elif section and section not in PROSE_SECTIONS:
            if _is_date_line(line):
                story.append(Paragraph(safe, date_style))
            else:
                head, detail = _split_entry_detail(line, section)
                if detail:
                    story.append(
                        Paragraph(
                            f"<b>{html.escape(head)}</b> - {html.escape(detail)}",
                            entry_detail_style,
                        )
                    )
                else:
                    story.append(Paragraph(safe, entry_style))
        else:
            story.append(Paragraph(safe, body_style))
        nonempty_seen += 1

    doc = SimpleDocTemplate(
        output, pagesize=A4, leftMargin=0.65 * inch, rightMargin=0.65 * inch,
        topMargin=0.55 * inch, bottomMargin=0.55 * inch,
        title=title, author="My_AGENT - Satyam Singh Bhadoriya",
    )
    doc.build(story)
    return output.getvalue()


def _shade_paragraph(paragraph, fill: str) -> None:
    """Fill a DOCX heading paragraph so the section stands out."""
    p_pr = paragraph._p.get_or_add_pPr()
    existing = p_pr.find(qn("w:shd"))
    if existing is not None:
        p_pr.remove(existing)
    shading = OxmlElement("w:shd")
    shading.set(qn("w:val"), "clear")
    shading.set(qn("w:color"), "auto")
    shading.set(qn("w:fill"), fill)
    p_pr.append(shading)


def _rgb(hex_colour: str) -> RGBColor:
    """Convert a six-character RGB hex colour into a python-docx colour."""
    return RGBColor.from_string(hex_colour)


def _rule_below_paragraph(paragraph, colour: str = "000000") -> None:
    """Draw a line under a DOCX heading, matching the reference CV."""
    p_pr = paragraph._p.get_or_add_pPr()
    existing = p_pr.find(qn("w:pBdr"))
    if existing is not None:
        p_pr.remove(existing)
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), colour)
    borders.append(bottom)
    p_pr.append(borders)


def _clear_table_borders(table) -> None:
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    if tbl_pr is None:
        tbl_pr = OxmlElement("w:tblPr")
        tbl.insert(0, tbl_pr)
    existing = tbl_pr.find(qn("w:tblBorders"))
    if existing is not None:
        tbl_pr.remove(existing)
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        item = OxmlElement(f"w:{edge}")
        item.set(qn("w:val"), "nil")
        item.set(qn("w:sz"), "0")
        item.set(qn("w:space"), "0")
        item.set(qn("w:color"), "auto")
        borders.append(item)
    tbl_pr.append(borders)


def build_ats_docx(
    text: str,
    title: str = "ATS Resume",
    photo_bytes: bytes | None = None,
    template: str = "reference",
) -> bytes:
    """Create a single-column DOCX preferred by many ATS platforms."""
    design = cv_template(template)
    ruled = design.get("ruled_headings", False)
    compact = design.get("compact", False)
    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.55)
    section.bottom_margin = Inches(0.55)
    section.left_margin = Inches(0.65)
    section.right_margin = Inches(0.65)

    normal = document.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(10)
    normal.paragraph_format.space_after = Pt(1 if compact else 2)
    normal.paragraph_format.line_spacing = 1.0 if compact else 1.05

    source_lines = text.splitlines()
    nonempty_seen = 0
    if photo_bytes:
        header, rest = _split_header_lines(text)
        table = document.add_table(rows=1, cols=2)
        _clear_table_borders(table)
        left_cell, right_cell = table.cell(0, 0), table.cell(0, 1)
        left_cell.width = Inches(5.8)
        right_cell.width = Inches(1.25)
        for index, line in enumerate(header):
            paragraph = (
                left_cell.paragraphs[0] if index == 0 else left_cell.add_paragraph()
            )
            run = paragraph.add_run(line)
            run.font.name = "Arial"
            if index == 0:
                run.bold = True
                run.font.size = Pt(18)
                run.font.color.rgb = _rgb(design["primary"])
            elif index == 1 and line == line.upper() and len(line.split()) <= 6:
                run.bold = True
                run.font.size = Pt(11.5)
                run.font.color.rgb = _rgb(design["accent"])
            else:
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(0x2A, 0x2E, 0x3A)
            paragraph.paragraph_format.space_after = Pt(3)
        photo_para = right_cell.paragraphs[0]
        photo_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        photo_para.add_run().add_picture(
            prepare_profile_photo(photo_bytes), width=Inches(1.12)
        )
        source_lines = rest
        nonempty_seen = 3

    section = ""
    for raw in source_lines:
        line = raw.strip()
        if not line:
            document.add_paragraph().paragraph_format.space_after = Pt(1)
            continue
        upper = line.upper().rstrip(":")
        if nonempty_seen == 0:
            paragraph = document.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = paragraph.add_run(line)
            run.bold = True
            run.font.name = "Arial"
            run.font.size = Pt(18)
            run.font.color.rgb = _rgb(design["primary"])
            paragraph.paragraph_format.space_after = Pt(3)
        elif nonempty_seen == 1 and line == upper and len(line.split()) <= 6:
            paragraph = document.add_paragraph()
            run = paragraph.add_run(line)
            run.bold = True
            run.font.name = "Arial"
            run.font.size = Pt(11.5)
            run.font.color.rgb = _rgb(design["accent"])
            paragraph.paragraph_format.space_after = Pt(3)
        elif upper in SECTION_HEADINGS or upper == "JD-MATCHED SKILLS":
            section = HEADING_ALIASES.get(upper, upper)
            paragraph = document.add_paragraph()
            run = paragraph.add_run(upper)
            run.bold = True
            run.font.name = "Arial"
            run.font.size = Pt(11)
            run.font.color.rgb = _rgb(design["heading_text"])
            paragraph.paragraph_format.space_before = Pt(7 if compact else 9)
            paragraph.paragraph_format.space_after = Pt(2 if compact else 5)
            if ruled:
                _rule_below_paragraph(paragraph)
            else:
                _shade_paragraph(paragraph, design["heading_fill"])
        elif line.startswith(("-", "*", "•")):
            paragraph = document.add_paragraph(style="List Bullet")
            paragraph.add_run(line.lstrip("-*• ").strip())
            paragraph.paragraph_format.space_after = Pt(1 if compact else 2)
        elif _is_numbered_item(line):
            # Keep the author's own numbering instead of letting Word renumber.
            paragraph = document.add_paragraph()
            paragraph.add_run(line)
            paragraph.paragraph_format.left_indent = Inches(0.25)
            paragraph.paragraph_format.space_after = Pt(1 if compact else 2)
        elif section and section not in PROSE_SECTIONS:
            paragraph = document.add_paragraph()
            if _is_date_line(line):
                run = paragraph.add_run(line)
                run.font.name = "Arial"
                run.font.size = Pt(9.5)
                run.italic = True
            else:
                head, detail = _split_entry_detail(line, section)
                run = paragraph.add_run(head)
                run.font.name = "Arial"
                run.font.size = Pt(10)
                run.bold = True
                if detail:
                    tail = paragraph.add_run(" - " + detail)
                    tail.font.name = "Arial"
                    tail.font.size = Pt(10)
            paragraph.paragraph_format.space_before = Pt(2 if compact else 3)
            paragraph.paragraph_format.space_after = Pt(0 if compact else 1)
        else:
            document.add_paragraph(line)
        nonempty_seen += 1

    document.core_properties.title = title
    document.core_properties.author = "My_AGENT - Satyam Singh Bhadoriya"
    output = BytesIO()
    document.save(output)
    return output.getvalue()


def build_cv_xlsx(
    text: str,
    title: str = "CV Reference",
    photo_bytes: bytes | None = None,
    template: str = "reference",
) -> bytes:
    """Create an Excel reference copy; XLSX is not recommended for ATS upload."""
    design = cv_template(template)
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "CV"
    sheet.column_dimensions["A"].width = 92
    sheet.column_dimensions["B"].width = 18
    sheet.sheet_view.showGridLines = False
    sheet.freeze_panes = "A2"

    if photo_bytes:
        xl_image = XLImage(prepare_profile_photo(photo_bytes))
        xl_image.width = 92
        xl_image.height = 114
        sheet.add_image(xl_image, "B1")

    row = 1
    nonempty_seen = 0
    for raw in text.splitlines():
        line = raw.strip()
        if not line:
            row += 1
            continue
        upper = line.upper().rstrip(":")
        cell = sheet.cell(row=row, column=1, value=line)
        cell.alignment = Alignment(wrap_text=True, vertical="top")
        if nonempty_seen == 0:
            cell.font = Font(
                name="Arial", size=17, bold=True, color=design["primary"]
            )
            sheet.row_dimensions[row].height = 27
        elif nonempty_seen == 1 and line == upper and len(line.split()) <= 6:
            cell.font = Font(
                name="Arial", size=11, bold=True, color=design["accent"]
            )
            sheet.row_dimensions[row].height = 20
        elif upper in SECTION_HEADINGS or upper == "JD-MATCHED SKILLS":
            cell.value = upper
            cell.font = Font(
                name="Arial", size=11, bold=True, color=design["heading_text"]
            )
            if design.get("ruled_headings"):
                cell.border = Border(bottom=Side(style="thin", color="000000"))
            else:
                cell.fill = PatternFill("solid", fgColor=design["heading_fill"])
            sheet.row_dimensions[row].height = 22
        else:
            cell.font = Font(name="Arial", size=10, color="2A2E3A")
            sheet.row_dimensions[row].height = 18
        nonempty_seen += 1
        row += 1

    sheet["C1"] = title
    sheet.column_dimensions["C"].hidden = True
    output = BytesIO()
    workbook.save(output)
    return output.getvalue()


def safe_output_stem(meta: dict) -> str:
    original = Path(meta.get("original_name", "CV.pdf")).stem
    base = re.sub(r"[^A-Za-z0-9]+", "_", original).strip("_") or "Tailored_CV"
    return f"{base}_JD_ATS"


def safe_output_name(meta: dict) -> str:
    return safe_output_stem(meta) + ".pdf"


def save_generated(pdf_bytes: bytes, filename: str) -> Path:
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    path = OUTPUT_DIR / filename
    path.write_bytes(pdf_bytes)
    return path
